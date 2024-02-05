#!/usr/bin/python2
"""
This module contains the docx functions and its utility functions
"""
from sys import argv
from docx import Document
from parse_txt import txt_to_dict


def parser(filename):
    """
    Parses `.docx` files to a uni-level dict object
    filename: the name/path of the .docx file to be parsed
    """
    doc = Document(filename)
    text = ""
    for p in doc.paragraphs:
        unicode = p.text.encode('ascii', "ignore")
        #decode byte to str
        if type(unicode) == bytes:
            unicode = unicode.decode('utf-8')
        text += unicode + '\n'


    # text = ("".join([p.text.encode("ascii", "ignore") #+ "\n"
    #                  for p in doc.paragraphs]))
    return txt_to_dict(text)


# execute if module is run
if __name__ == "__main__":
    if len(argv) > 1:
        filename = argv[1]
        print(parser(filename))
    else:
        print("usage: {} [filename]".format(argv[0]))
